from __future__ import annotations

from collections.abc import Sequence
from io import BytesIO

import docx  # type: ignore[import-untyped]
import pytesseract
from PIL import Image
from PyPDF2 import PdfReader
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from backend.app.models.db.documents import Document, DocumentChunk
from backend.app.services.embeddings import EmbeddingService


class DocumentService:
    """Persistence helper for documents and their vector representations."""

    def __init__(self, session: AsyncSession) -> None:
        self.session = session

    async def ingest_text(
        self,
        *,
        title: str,
        source: str | None,
        content: str,
        meta: dict | None = None,
        chunk_size: int = 800,
        chunk_overlap: int = 80,
        embed: bool = True,
        embedding_service: EmbeddingService | None = None,
    ) -> Document:
        chunks = self._chunk_text(content, chunk_size=chunk_size, overlap=chunk_overlap)
        if not chunks:
            raise ValueError("Document must contain readable text.")

        embeddings: Sequence[Sequence[float]] | None = None
        if embed:
            service = embedding_service or EmbeddingService()
            embeddings = await service.embed(chunks)

        return await self.create_document(
            title=title,
            source=source,
            meta=meta,
            chunks=chunks,
            embeddings=embeddings,
        )

    async def create_document(
        self,
        *,
        title: str,
        source: str | None = None,
        meta: dict | None = None,
        chunks: Sequence[str],
        embeddings: Sequence[Sequence[float]] | None = None,
    ) -> Document:
        document = Document(title=title, source=source, meta=meta or {})

        for idx, chunk_text in enumerate(chunks):
            embedding = None
            if embeddings is not None and idx < len(embeddings):
                embedding = list(embeddings[idx])

            document.chunks.append(
                DocumentChunk(chunk_index=idx, content=chunk_text, embedding=embedding)
            )

        self.session.add(document)
        await self.session.flush()
        await self.session.commit()
        return document

    async def ingest_file(
        self,
        *,
        filename: str,
        content_type: str | None,
        data: bytes,
        title: str | None = None,
        source: str | None = None,
        meta: dict | None = None,
        chunk_size: int = 800,
        chunk_overlap: int = 80,
        embed: bool = True,
        embedding_service: EmbeddingService | None = None,
    ) -> Document:
        """Ingest an arbitrary uploaded file (PDF, DOCX, text, image with OCR)."""

        text = self._extract_text_from_file(filename=filename, content_type=content_type, data=data)
        if not text.strip():
            raise ValueError("Uploaded file is empty or unreadable.")

        inferred_title = title or filename or "Untitled"
        inferred_source = source or filename
        meta = dict(meta or {})
        meta.setdefault("filename", filename)
        if content_type:
            meta.setdefault("content_type", content_type)

        return await self.ingest_text(
            title=inferred_title,
            source=inferred_source,
            content=text,
            meta=meta,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            embed=embed,
            embedding_service=embedding_service,
        )

    async def list_documents(self) -> list[Document]:
        stmt = select(Document).options(selectinload(Document.chunks))
        result = await self.session.execute(stmt)
        return list(result.scalars())

    @staticmethod
    def _chunk_text(content: str, *, chunk_size: int = 800, overlap: int = 80) -> list[str]:
        normalized = content.replace("\r\n", "\n").strip()
        if not normalized:
            return []

        overlap = max(0, min(overlap, chunk_size // 2))
        chunks: list[str] = []
        start = 0
        text_length = len(normalized)

        while start < text_length:
            end = min(text_length, start + chunk_size)
            chunk = normalized[start:end].strip()
            if chunk:
                chunks.append(chunk)
            if end == text_length:
                break
            start = max(end - overlap, start + 1)

        return chunks

    @staticmethod
    def _extract_text_from_file(
        *, filename: str, content_type: str | None, data: bytes
    ) -> str:
        """Best-effort text extraction for PDF, DOCX, plaintext, and images (OCR)."""

        if not data:
            return ""

        name_lower = (filename or "").lower()
        content_type = (content_type or "").lower()

        # PDF
        if name_lower.endswith(".pdf") or "pdf" in content_type:
            try:
                reader = PdfReader(BytesIO(data))
                parts: list[str] = []
                for page in reader.pages:
                    page_text = page.extract_text() or ""
                    parts.append(page_text)
                return "\n".join(parts)
            except Exception:
                return ""

        # DOCX
        if name_lower.endswith(".docx") or "word" in content_type:
            try:
                document = docx.Document(BytesIO(data))
                return "\n".join(paragraph.text for paragraph in document.paragraphs)
            except Exception:
                return ""

        # Images (basic OCR)
        if any(
            name_lower.endswith(ext)
            for ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp")
        ) or content_type.startswith("image/"):
            try:
                image = Image.open(BytesIO(data))
                return pytesseract.image_to_string(image)
            except Exception:
                return ""

        # Fallback: try text decode
        for encoding in ("utf-8", "latin-1"):
            try:
                return data.decode(encoding)
            except UnicodeDecodeError:
                continue

        return data.decode("utf-8", errors="ignore")
