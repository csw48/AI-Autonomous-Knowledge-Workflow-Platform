"""Generate sample DOCX and PDF documents for upload testing."""

from __future__ import annotations

from pathlib import Path

from docx import Document  # type: ignore[import-untyped]
from reportlab.pdfgen import canvas  # type: ignore[import-untyped]


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    samples_dir = root / "sample_docs"
    samples_dir.mkdir(exist_ok=True)

    # DOCX
    doc = Document()
    doc.add_heading("Sample DOCX for RAG upload", level=1)
    doc.add_paragraph(
        "This is a sample DOCX document used to test the upload and RAG pipeline."
    )
    doc.add_paragraph(
        "It contains a few sentences about hello world, PDF and DOCX processing "
        "in our AI Autonomous Knowledge & Workflow Platform."
    )
    docx_path = samples_dir / "sample-notes.docx"
    doc.save(docx_path)

    # PDF
    pdf_path = samples_dir / "sample-notes.pdf"
    c = canvas.Canvas(str(pdf_path))
    c.setTitle("Sample PDF for RAG upload")
    text = c.beginText(72, 720)
    text.textLine("Sample PDF for RAG upload")
    text.textLine(
        "This PDF contains text about hello world, PDF indexing, and RAG search."
    )
    text.textLine(
        "You can upload it through the dashboard to verify PDF extraction."
    )
    c.drawText(text)
    c.showPage()
    c.save()

    print(f"Created: {docx_path}")
    print(f"Created: {pdf_path}")


if __name__ == "__main__":
    main()

